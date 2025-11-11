from flask import Flask, request, jsonify, Response
from datetime import datetime
from docx import Document
import re, os, traceback, yagmail
import threading

app = Flask(__name__)

import logging
app.logger.setLevel(logging.INFO)

# ----- CONFIG -----
GMAIL_USER = os.getenv("GMAIL_USER")
GMAIL_PASS = os.getenv("GMAIL_PASS")

# Strong regex parser for messages like:
# "quote 110 for Raju at Raj Pvt Ltd, 500 pcs 3in SS 316L sheets at 25000, hsn 7219, email raju@example.com"
QUOTE_RE = re.compile(
    r"""
    quote\s*(?P<qno>\d+)                              # quote number
    .*?\bfor\b\s+(?P<name>[A-Za-z][A-Za-z ]{1,60})    # name
    \s+\bat\b\s+(?P<company>[^,]+)                    # company
    ,?\s+(?P<qty>\d+)\s*(?P<units>pcs|nos|kgs|kg|mt|ton|piece|pieces)? # quantity + optional units
    .*?\b(?P<product>ss|stainless|mild|ms|alloy|aluminium|copper|brass|nickel|3in|4in|pipe|sheet|sheets|coil|bar|rod|flange|valve|fitting|316L|304L|310S|duplex|superduplex|sch\s*\d+|sch\d+|[\w\-\s/]+?)\b
    .*?\bat\s+(?P<rate>\d{2,})                        # rate
    (?:.*?\bhsn\b\s*(?P<hsn>\d{4,8}))?                # optional hsn
    .*?\bemail\b\s*(?P<email>[\w.\-\+%]+@[\w.-]+\.[A-Za-z]{2,})  # email
    """,
    re.IGNORECASE | re.VERBOSE | re.DOTALL,
)

def _clean(s: str) -> str:
    return " ".join((s or "").strip().split())

def parse_message(text: str):
    if not text:
        return None

    raw = text
    text = " ".join(text.split())
    app.logger.info(f"[PARSE] incoming: {raw}")

    import re

    # ---------- helpers & regex ----------
    EMAIL_RE = re.compile(r'([\w.\-+%]+@[\w.\-]+\.[A-Za-z]{2,})')
    QNO_RE   = re.compile(r'\bquote\s*(\d{1,10})\b', re.I)
    RATE_RE  = re.compile(r'\b(?:rate|at)\s*([0-9]{3,})\b', re.I)
    # allow "5pcs", "5 psc", "5psc", "5 nos" etc
    QTY_RE   = re.compile(r'\b(\d{1,7})\s*(pcs|psc|nos|pieces?|kgs?|kg|mt|ton|bundle|bndl)?\b', re.I)
    HSN_RE   = re.compile(r'\bhsn\s*([0-9]{4,8})\b', re.I)

    units_map = {
        "piece":"pcs","pieces":"pcs","nos":"pcs","pcs":"pcs","psc":"pcs",
        "kg":"Kgs","kgs":"Kgs","mt":"MT","ton":"Ton","bundle":"Bundle","bndl":"Bundle"
    }

    # Name patterns: "for NAME", "to NAME", "customer name is NAME"
    NAME_PATTERNS = [
        re.compile(r'\bfor\s+([A-Za-z][\w .\'\-]{1,60})\b', re.I),
        re.compile(r'\bto\s+([A-Za-z][\w .\'\-]{1,60})\b', re.I),
        re.compile(r'\bcustomer\s+name\s+is\s+([A-Za-z][\w .\'\-]{1,60})\b', re.I),
    ]

    # Company optional: "at COMPANY" if it exists
    COMPANY_RE = re.compile(r'\bat\s+([^,]+)', re.I)

    # ---------- try structured patterns first ----------
    patterns = [
        re.compile(
            r"""quote\s*(?P<qno>\d+).*?\b(?:for|to)\b\s+(?P<name>[^,]+?)
                (?:\s+\bat\b\s+(?P<company>[^,]+))?
                .*?(?P<qty>\d{1,7})\s*(?P<units>pcs|psc|nos|pieces?|kgs?|kg|mt|ton|bundle|bndl)?
                \s+(?P<product>.+?)\s+\b(?:rate|at)\b\s+(?P<rate>\d{3,})
                (?:.*?\bhsn\b\s*(?P<hsn>\d{4,8}))?
                .*?\bemail\b\s*(?P<email>[\w.\-+%]+@[\w.\-]+\.[A-Za-z]{2,})
            """, re.I | re.X),
    ]

    def norm_units(u):
        u = (u or "").lower()
        return units_map.get(u, "pcs")

    for pat in patterns:
        m = pat.search(text)
        if m:
            d = {k: _clean(v) if isinstance(v, str) else v for k, v in m.groupdict().items()}
            return {
                "qno": d.get("qno",""),
                "name": (d.get("name","") or "").title(),
                "company": d.get("company",""),
                "qty": d.get("qty",""),
                "units": norm_units(d.get("units")) if d.get("qty") else "",
                "product": (d.get("product","") or "")[:120],
                "rate": d.get("rate",""),
                "hsn": d.get("hsn",""),
                "email": d.get("email",""),
            }

    # ---------- heuristic fallback ----------
    qno    = (QNO_RE.search(text) or [None, ""])[1]
    email  = (EMAIL_RE.search(text) or [None, ""])[1]
    rate   = (RATE_RE.search(text) or [None, ""])[1]

    qtym   = None
    # prefer the 'quantity/qty' segment if present; else any number+unit
    qty_kw = re.search(r'\b(quantity|qty)\b\s*(\d{1,7}\s*(?:pcs|psc|nos|pieces?|kgs?|kg|mt|ton|bundle|bndl)?)', text, re.I)
    if qty_kw:
        qtym = QTY_RE.search(qty_kw.group(0))
    if not qtym:
        qtym = QTY_RE.search(text)

    qty    = qtym.group(1) if qtym else ""
    uraw   = qtym.group(2) if (qtym and qtym.lastindex and qtym.lastindex >= 2) else ""
    units  = norm_units(uraw) if qty else ""
    hsn    = (HSN_RE.search(text) or [None, ""])[1]

    # name: try multiple patterns
    name = ""
    for np in NAME_PATTERNS:
        mm = np.search(text)
        if mm:
            name = _clean(mm.group(1))
            break

    # company (optional)
    company = ""
    mc = COMPANY_RE.search(text)
    if mc:
        company = _clean(mc.group(1))

    # product: take the segment between the name and the first of (quantity|qty|rate|hsn|email)
    product = ""
    anchors = []
    for a in [" quantity ", " qty ", " rate ", " hsn ", " email "]:
        p = text.lower().find(a)
        if p != -1:
            anchors.append(p)
    cut_to = min(anchors) if anchors else -1

    # starting point: after name, or after "to/for/customer name is"
    start = -1
    if name:
        # find where that name appears and begin after its occurrence
        for tag in [" for ", " to ", " customer name is "]:
            pos = text.lower().find(tag)
            if pos != -1:
                pos2 = pos + len(tag)
                # next substring begins with the name
                name_pos = text.lower().find(name.lower(), pos2)
                if name_pos != -1:
                    start = name_pos + len(name)
                    break
    # fallback: take from beginning
    if start == -1:
        start = 0
    if cut_to != -1 and cut_to > start:
        candidate = text[start:cut_to]
    else:
        candidate = text[start:]

    # strip commas and filler words
    candidate = candidate.strip(" ,.-")
    # if candidate begins with delimiters like 'customer name is', trim again
    candidate = re.sub(r'^(customer name is|for|to)\b', '', candidate, flags=re.I).strip(" ,.-")

    # remove any trailing 'at <rate>' part if present
    if rate:
        candidate = re.sub(fr'\b(?:rate|at)\s*{re.escape(rate)}\b.*$', '', candidate, flags=re.I).strip(" ,.-")

    # If it still looks empty but we had a "quantity ...", try the text between the start of the line and "quantity"
    if not candidate and qty_kw:
        head = text[: text.lower().find(qty_kw.group(1).lower())]
        candidate = head.strip(" ,.-")

    product = candidate[:120]

    ctx = {
        "qno": qno,
        "name": name.title() if name else "",
        "company": company,          # optional
        "qty": qty,
        "units": units if qty else "",
        "product": product,
        "rate": rate,
        "hsn": hsn,
        "email": email,
    }

    # Only require name, qty, rate, email; company is optional
    required = ["name", "qty", "rate", "email"]
    if any(not ctx[k] for k in required):
        app.logger.warning(f"[PARSE] fallback incomplete -> {ctx}")
        return None

    return ctx

def create_doc(ctx:dict) -> str:
    doc = Document()
    doc.add_heading(f'Quotation #{ctx["qno"]}', level=1)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%d-%b-%Y")}')
    doc.add_paragraph(f'Customer: {ctx["name"]}')
    doc.add_paragraph(f'Company: {ctx["company"]}')
    doc.add_paragraph(f'Product: {ctx["product"]}')
    doc.add_paragraph(f'Quantity: {ctx["qty"]} {ctx["units"]}')
    doc.add_paragraph(f'Rate: {ctx["rate"]}')
    if ctx.get("hsn"):
        doc.add_paragraph(f'HSN: {ctx["hsn"]}')
    doc.add_paragraph(f'Email: {ctx["email"]}')
    fname = f'Quotation_{ctx["name"].replace(" ","_")}_{datetime.now().date()}.docx'
    doc.save(fname)
    return fname

def send_email(attachment_path:str, to_email:str):
    if not GMAIL_USER or not GMAIL_PASS:
        app.logger.error("GMAIL_USER/GMAIL_PASS missing; skipping email send.")
        return False
    try:
        yag = yagmail.SMTP(GMAIL_USER, GMAIL_PASS)
        yag.send(
            to=to_email,
            subject="Quotation from Nivee Metal Products Pvt. Ltd.",
            contents="Please find the attached quotation.",
            attachments=attachment_path,
        )
        return True
    except Exception as e:
        app.logger.exception(f"Email send failed: {e}")
        return False


def _background_worker(text: str):
    """Background worker: parse, create doc, send email. Runs in a daemon thread."""
    try:
        app.logger.info(f"[BG] start processing: {text}")
        ctx = parse_message(text)
        if not ctx:
            app.logger.warning("[BG] parse failed; skipping")
            return
        file_path = create_doc(ctx)
        ok = send_email(file_path, ctx.get("email"))
        app.logger.info(f"[BG] done: emailed={ok} file={file_path}")
    except Exception as e:
        app.logger.exception(f"Exception in background worker: {e}")

def extract_text_from_meta(payload:dict) -> str | None:
    try:
        entry = payload.get("entry", [])[0]
        change = entry.get("changes", [])[0]
        value = change.get("value", {})
        msgs = value.get("messages", [])
        if not msgs:
            return None
        msg = msgs[0]
        if msg.get("type") == "text":
            return msg["text"]["body"]
        return None
    except Exception:
        return None

@app.route("/webhook", methods=["GET", "POST"])
def webhook():
    try:
        if request.method == "GET":
            # Support initial verification if you use this endpoint for Meta verification
            hub_mode = request.args.get("hub.mode")
            hub_challenge = request.args.get("hub.challenge")
            hub_verify_token = request.args.get("hub.verify_token")
            # Optional: compare to your env VERIFY_TOKEN
            if hub_mode == "subscribe" and hub_challenge:
                return Response(hub_challenge, status=200)
            return Response("OK", status=200)

        # POST
        body = request.get_json(silent=True) or {}
        # Accept both tester JSON and Meta payload
        app.logger.info(f"[WEBHOOK] body: {body}")
        text = body.get("message") or extract_text_from_meta(body)
        if not text:
            app.logger.info("Webhook received but no parsable text message; returning 200.")
            return jsonify({"status": "ignored"}), 200

        # Immediately acknowledge to WhatsApp / Meta to avoid retries
        try:
            thread = threading.Thread(target=_background_worker, args=(text,), daemon=True)
            thread.start()
            app.logger.info("[WEBHOOK] background worker started")
        except Exception as e:
            app.logger.exception(f"Failed to start background worker: {e}")

    # Return 200 "ok" immediately to acknowledge WhatsApp/Meta and avoid retries
    return Response("ok", status=200)

    except Exception as e:
        app.logger.error("Exception in /webhook: %s\n%s", e, traceback.format_exc())
        # Always 200 for Meta to avoid retry storms; log the error
        return jsonify({"status": "error_logged"}), 200

@app.route("/")
def health():
    return "Quotation bot is running (regex-only)."

if __name__ == "__main__":
    port = int(os.getenv("PORT", "10000"))
    app.run(host="0.0.0.0", port=port)
